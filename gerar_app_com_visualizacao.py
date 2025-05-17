import streamlit as st
import pandas as pd
import zipfile
import io
from datetime import datetime
from docx import Document

st.set_page_config(page_title="GeraR – EC Serviços", layout="centered")

st.title("GeraR – Gerador de Relatórios por Localidade")
st.write("Envie suas planilhas mensais e o modelo Word para gerar os relatórios em PDF por localidade.")

uploaded_modelo = st.file_uploader("Modelo Word (.docx)", type="docx")
uploaded_planilhas = st.file_uploader("Planilhas Excel (.xlsx)", type="xlsx", accept_multiple_files=True)

meses_disponiveis = ["jan", "fev", "mar", "abr", "mai", "jun", "jul", "ago", "set", "out", "nov", "dez"]
meses_selecionados = st.multiselect("Selecionar meses a considerar", meses_disponiveis, default=["jan", "fev", "mar"])

def normalizar_texto(txt):
    return str(txt).strip().lower() if pd.notna(txt) else ""

def primeira_maiuscula(txt):
    return txt.capitalize() if isinstance(txt, str) else ""

def substituir_placeholder_em_paragrafos(paragraphs, marcador, valor):
    for paragraph in paragraphs:
        texto_completo = ''.join(run.text for run in paragraph.runs)
        if marcador in texto_completo:
            novo_texto = texto_completo.replace(marcador, valor)
            for run in paragraph.runs:
                run.text = ""
            paragraph.runs[0].text = novo_texto

def consolidar_dados(planilhas):
    consolidado = []
    ignoradas = set()
    for i, file in enumerate(planilhas):
        df = pd.read_excel(file)
        mes = meses_disponiveis[i] if i < len(meses_disponiveis) else f"mes{i+1}"
        for _, row in df.iterrows():
            loc = normalizar_texto(row.iloc[3])
            grp = normalizar_texto(row.iloc[8])
            itm = normalizar_texto(row.iloc[11])
            obs = normalizar_texto(row.iloc[27])
            if not itm or not loc:
                ignoradas.add(loc if loc else "localidade indefinida")
                continue
            chave = (loc, itm, grp)
            existente = next((x for x in consolidado if x["chave"] == chave), None)
            if existente:
                existente["meses"].add(mes)
                if obs and (not existente["obs"] or len(obs) > len(existente["obs"])):
                    existente["obs"] = obs
            else:
                consolidado.append({
                    "chave": chave,
                    "localidade": loc,
                    "item": itm,
                    "grupo": grp,
                    "obs": obs,
                    "meses": {mes}
                })
    return consolidado, ignoradas

if uploaded_modelo and uploaded_planilhas:
    dados, ignoradas = consolidar_dados(uploaded_planilhas)

    if dados:
        st.subheader("Pré-visualização dos dados consolidados")
        preview = pd.DataFrame([{
            "Localidade": d["localidade"].title(),
            "Item": primeira_maiuscula(d["item"]),
            "Grupo": primeira_maiuscula(d["grupo"]),
            "Observação": primeira_maiuscula(d["obs"]),
            "Meses": ", ".join(sorted(d["meses"]))
        } for d in dados])
        st.dataframe(preview)

        if ignoradas:
            st.warning("Localidades ignoradas (sem dados válidos): " + ", ".join(sorted(ignoradas)))

        if st.button("Gerar documentos"):
            agrupado = {}
            for d in dados:
                agrupado.setdefault(d["localidade"], []).append(d)

            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, "w") as zf:
                for loc, itens in agrupado.items():
                    doc = Document(uploaded_modelo)
                    substituir_placeholder_em_paragrafos(doc.paragraphs, "<<localidade>>", loc.title())
                    substituir_placeholder_em_paragrafos(doc.paragraphs, "<<datadaemissao>>", datetime.now().strftime("%d de %B de %Y"))

                    tabela = doc.tables[0]
                    base = tabela.rows[2]
                    tabela._tbl.remove(base._tr)
                    for i, item in enumerate(itens, 1):
                        nova = tabela.add_row().cells
                        nova[0].text = primeira_maiuscula(item["item"])
                        nova[1].text = primeira_maiuscula(item["obs"])
                        nova[2].text = primeira_maiuscula(item["grupo"])
                        for j, mes in enumerate(["jan", "fev", "mar"]):
                            if mes in item["meses"]:
                                nova[3 + j].text = "X"

                    filename = f"{loc.replace(' ', '_')}_{'_'.join(meses_selecionados)}.docx"
                    file_stream = io.BytesIO()
                    doc.save(file_stream)
                    zf.writestr(filename, file_stream.getvalue())

            st.success("Documentos gerados com sucesso!")
            st.download_button("Baixar ZIP com os relatórios", data=zip_buffer.getvalue(), file_name="relatorios.zip")
    else:
        st.info("Nenhum dado válido encontrado para gerar relatório.")
