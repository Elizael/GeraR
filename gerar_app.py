import streamlit as st
import pandas as pd
import zipfile
import io
from datetime import datetime
from docx import Document
import os

st.set_page_config(page_title="GeraR – EC Serviços", layout="centered")

st.title("GeraR – Gerador de Relatórios por Localidade")
st.write("Envie suas planilhas mensais e o modelo Word para gerar os relatórios em PDF por localidade.")

uploaded_modelo = st.file_uploader("Modelo Word (.docx)", type="docx")
uploaded_planilhas = st.file_uploader("Planilhas Excel (.xlsx)", type="xlsx", accept_multiple_files=True)

meses_disponiveis = ["jan", "fev", "mar", "abr", "mai", "jun", "jul", "ago", "set", "out", "nov", "dez"]
meses_selecionados = st.multiselect("Selecionar meses a considerar", meses_disponiveis, default=["jan", "fev", "mar"])

def normalizar_texto(txt):
    return str(txt).strip().lower() if pd.notna(txt) else ""

def primeira_maiuscula(txt):
    return txt.capitalize() if isinstance(txt, str) else ""

def consolidar_dados(planilhas):
    consolidado = []
    for i, file in enumerate(planilhas):
        df = pd.read_excel(file)
        mes = meses_disponiveis[i] if i < len(meses_disponiveis) else f"mes{i+1}"
        for _, row in df.iterrows():
            loc = normalizar_texto(row.iloc[3])
            grp = normalizar_texto(row.iloc[8])
            itm = normalizar_texto(row.iloc[11])
            obs = normalizar_texto(row.iloc[27])
            if not itm or not loc:
                continue
            chave = (loc, itm, grp)
            existente = next((x for x in consolidado if x["chave"] == chave), None)
            if existente:
                existente["meses"].add(mes)
                if obs and (not existente["obs"] or len(obs) > len(existente["obs"])):
                    existente["obs"] = obs
            else:
                consolidado.append({
                    "chave": chave,
                    "localidade": loc,
                    "item": itm,
                    "grupo": grp,
                    "obs": obs,
                    "meses": {mes}
                })
    return consolidado

if uploaded_modelo and uploaded_planilhas:
    if st.button("Gerar documentos"):
        dados = consolidar_dados(uploaded_planilhas)
        agrupado = {}
        for d in dados:
            agrupado.setdefault(d["localidade"], []).append(d)

        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "w") as zf:
            for loc, itens in agrupado.items():
                doc = Document(uploaded_modelo)
                for p in doc.paragraphs:
                    if "<<localidade>>" in p.text:
                        p.text = p.text.replace("<<localidade>>", loc.title())
                    if "<<datadaemissao>>" in p.text:
                        hoje = datetime.now().strftime("%d de %B de %Y")
                        p.text = p.text.replace("<<datadaemissao>>", hoje)

                tabela = doc.tables[0]
                base = tabela.rows[2]
                tabela._tbl.remove(base._tr)
                for i, item in enumerate(itens, 1):
                    nova = tabela.add_row().cells
                    nova[0].text = primeira_maiuscula(item["item"])
                    nova[1].text = primeira_maiuscula(item["obs"])
                    nova[2].text = primeira_maiuscula(item["grupo"])
                    for j, mes in enumerate(["jan", "fev", "mar"]):
                        if mes in item["meses"]:
                            nova[3 + j].text = "X"

                filename = f"{loc.replace(' ', '_')}_{'_'.join(meses_selecionados)}.docx"
                file_stream = io.BytesIO()
                doc.save(file_stream)
                zf.writestr(filename, file_stream.getvalue())

        st.success("Documentos gerados com sucesso!")
        st.download_button("Baixar ZIP com os relatórios", data=zip_buffer.getvalue(), file_name="relatorios.zip")
